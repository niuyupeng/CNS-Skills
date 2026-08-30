#!/usr/bin/env python3
"""Build the public, fully synthetic CNS Skills before/after DOCX demo."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "CNS-Skills-synthetic-manuscript-demo.unscrubbed.docx"
BEFORE = ROOT / "input_manuscript.md"
AFTER = ROOT / "revised_manuscript.md"
FIGURE = ROOT / "figure-1.png"

NAVY = "102A43"
INK = "243B53"
TEAL = "0F766E"
TEAL_LIGHT = "E6F4F1"
AMBER = "D97706"
AMBER_LIGHT = "FFF4DD"
RED = "B42318"
RED_LIGHT = "FDECEC"
BLUE_LIGHT = "EAF2F8"
GRAY = "627D98"
GRAY_LIGHT = "F4F7FA"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color: str, size: str = "8", sides=("top", "left", "bottom", "right")) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side in sides:
        tag = f"w:{side}"
        edge = borders.find(qn(tag))
        if edge is None:
            edge = OxmlElement(tag)
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), size)
        edge.set(qn("w:color"), color)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, end))


def style_run(run, *, size=None, bold=None, color=None, italic=None, font="Aptos") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(paragraph, *, before=0, after=5, line=1.08) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_together = False
    fmt.widow_control = True


def add_text(document, text: str, *, style=None, size=9.4, color=INK, bold=False, italic=False, after=5):
    paragraph = document.add_paragraph(style=style)
    style_paragraph(paragraph, after=after)
    run = paragraph.add_run(text)
    style_run(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_kicker(document, text: str, color=TEAL) -> None:
    paragraph = document.add_paragraph()
    style_paragraph(paragraph, before=2, after=3)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text.upper())
    style_run(run, size=8.2, color=color, bold=True)
    run.font.all_caps = True
    run.font.spacing = Pt(0.8)


def add_heading(document, text: str, level=1, color=NAVY) -> None:
    paragraph = document.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    style_run(run, color=color)


def add_page_break(document) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def markdown_title(path: Path) -> str:
    for line in path.read_text(encoding="utf-8").splitlines():
        if line.startswith("# "):
            return line[2:].strip()
    raise ValueError(f"missing title in {path}")


def markdown_section(path: Path, name: str) -> str:
    lines = path.read_text(encoding="utf-8").splitlines()
    start = None
    selected: list[str] = []
    for index, line in enumerate(lines):
        if line.strip() == f"## {name}":
            start = index + 1
            continue
        if start is not None and index >= start:
            if line.startswith("## "):
                break
            if line.startswith("### "):
                selected.append(line[4:].strip() + ".")
                continue
            if line.startswith("|") or re.match(r"^Table 1\.", line):
                continue
            if line.startswith("Figure 1."):
                continue
            if line.startswith(">"):
                continue
            if line.strip():
                selected.append(re.sub(r"`([^`]+)`", r"\1", line.strip()))
    return "\n\n".join(selected)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)
    section.different_first_page_header_footer = True

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(9.4)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for level, size, before, after in ((1, 22, 10, 7), (2, 14, 8, 4), (3, 10.5, 5, 2)):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Aptos Display" if level < 3 else "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), style.font.name)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(NAVY)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = document.styles["Caption"]
    caption.font.name = "Aptos"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    caption.font.size = Pt(8)
    caption.font.color.rgb = RGBColor.from_string(GRAY)
    caption.font.italic = False
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(6)

    if "Demo Label" not in document.styles:
        label = document.styles.add_style("Demo Label", WD_STYLE_TYPE.PARAGRAPH)
        label.font.name = "Aptos"
        label._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
        label.font.size = Pt(8)
        label.font.bold = True
        label.font.color.rgb = RGBColor.from_string(RED)
        label.paragraph_format.space_after = Pt(4)

    settings = document.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    props = document.core_properties
    props.title = "CNS Skills fully synthetic manuscript before/after demo"
    props.subject = "Evidence-bounded scientific manuscript revision demonstration"
    props.author = "CNS Skills"
    props.last_modified_by = "CNS Skills"
    props.keywords = "synthetic demonstration, scientific writing, evidence boundary"
    props.comments = "No real research, citations, participants, animals, or personal data."


def configure_running_matter(document: Document) -> None:
    section = document.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    table = header.add_table(rows=1, cols=2, width=Inches(7.05))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(4.95)
    table.columns[1].width = Inches(2.10)
    left, right = table.rows[0].cells
    left.text = "CNS SKILLS · SYNTHETIC MANUSCRIPT DEMO"
    right.text = "NO REAL CITATIONS"
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for cell in (left, right):
        set_cell_margins(cell, top=20, bottom=35, start=0, end=0)
        set_cell_border(cell, TEAL, size="12", sides=("bottom",))
        for run in cell.paragraphs[0].runs:
            style_run(run, size=7.2, bold=True, color=TEAL)

    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(paragraph, before=0, after=0)
    run = paragraph.add_run("FULLY SYNTHETIC · NOT SCIENTIFIC EVIDENCE   |   ")
    style_run(run, size=7.2, color=GRAY, bold=True)
    add_field(paragraph, "PAGE")
    for item in paragraph.runs[1:]:
        style_run(item, size=7.2, color=GRAY)


def add_banner(document, label: str, title: str, subtitle: str, fill=NAVY) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(7.0)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=170, bottom=170, start=220, end=220)
    paragraph = cell.paragraphs[0]
    style_paragraph(paragraph, after=5)
    run = paragraph.add_run(label.upper())
    style_run(run, size=8, bold=True, color="6EE7D2")
    paragraph = cell.add_paragraph()
    style_paragraph(paragraph, after=5, line=1.0)
    run = paragraph.add_run(title)
    style_run(run, size=23, bold=True, color=WHITE, font="Aptos Display")
    paragraph = cell.add_paragraph()
    style_paragraph(paragraph, after=0)
    run = paragraph.add_run(subtitle)
    style_run(run, size=10, color="D9E2EC")


def add_metric_strip(document) -> None:
    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    values = (("18", "toy rows"), ("12 / 6", "train / holdout"), ("1", "fictional batch"), ("0", "real citations"))
    for index, (value, label) in enumerate(values):
        cell = table.cell(0, index)
        cell.width = Inches(1.73)
        set_cell_shading(cell, GRAY_LIGHT if index % 2 == 0 else BLUE_LIGHT)
        set_cell_margins(cell, top=100, bottom=95, start=70, end=70)
        set_cell_border(cell, WHITE, size="12")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, after=0)
        r = p.add_run(value + "\n")
        style_run(r, size=18, bold=True, color=NAVY)
        r = p.add_run(label.upper())
        style_run(r, size=7, bold=True, color=GRAY)


def add_comparison_table(document) -> None:
    rows = [
        ("Workflow", "Autonomous closed loop", "Retrospective ranking only"),
        ("Evidence unit", "Three wells prove reproducibility", "Technical wells from one fictional batch"),
        ("Biological claim", "Universally safe and effective", "No safety, efficacy, or mechanism inference"),
        ("Translation", "Immediate clinical readiness", "Prospective independent validation still required"),
    ]
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (1.35, 2.75, 2.95)
    for index, (label, fill) in enumerate((("Decision", NAVY), ("Before", RED), ("After", TEAL))):
        cell = table.rows[0].cells[index]
        cell.width = Inches(widths[index])
        set_cell_shading(cell, fill)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        r = p.add_run(label)
        style_run(r, size=8.4, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for col, value in enumerate(values):
            cells[col].width = Inches(widths[col])
            set_cell_margins(cells[col], top=90, bottom=90)
            set_cell_shading(cells[col], WHITE if row_index % 2 == 0 else GRAY_LIGHT)
            set_cell_border(cells[col], "D9E2EC", size="6")
            p = cells[col].paragraphs[0]
            r = p.add_run(value)
            style_run(r, size=8.3, bold=(col == 0), color=INK)


def add_risk_table(document) -> None:
    rows = [
        ("Evidence boundary", "Universal safety / efficacy", "One-batch descriptive fixture", "Narrowed"),
        ("Autonomy", "Closed-loop discovery", "No next experiment was selected", "Reclassified"),
        ("Citations", "Missing 3; unused 2; duplicate DOI", "3/3 placed; all synthetic placeholders", "Resolved"),
        ("Displays", "Figure 2 missing; captions orphaned", "Figure 1 and Table 1 introduced in prose", "Resolved"),
        ("Clean copy", "TODO, editor note, author query, output label", "Reader-visible manuscript only", "Resolved"),
    ]
    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (1.28, 2.10, 2.50, 0.98)
    for index, label in enumerate(("Gate", "Before risk", "Evidence-bounded correction", "State")):
        cell = table.rows[0].cells[index]
        cell.width = Inches(widths[index])
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        r = cell.paragraphs[0].add_run(label)
        style_run(r, size=8, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for col, value in enumerate(values):
            cells[col].width = Inches(widths[col])
            set_cell_margins(cells[col], top=95, bottom=95)
            set_cell_shading(cells[col], WHITE if row_index % 2 == 0 else GRAY_LIGHT)
            set_cell_border(cells[col], "D9E2EC", size="5")
            p = cells[col].paragraphs[0]
            if col == 3:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(value)
            style_run(r, size=7.9, bold=(col in (0, 3)), color=TEAL if col == 3 else INK)


def add_manuscript_paragraphs(document, text: str) -> None:
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.endswith(".") and len(block) < 55 and block[0].isupper():
            add_heading(document, block[:-1], level=3, color=TEAL)
        else:
            add_text(document, block, size=9.2, after=5)


def add_results_table(document) -> None:
    table = document.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    headers = ("Formulation", "Pred. reduction", "Pred. viability", "Frozen reduction", "Frozen viability", "Permitted interpretation")
    widths = (0.80, 0.84, 0.84, 0.84, 0.84, 2.75)
    for index, label in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = Inches(widths[index])
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell, top=80, bottom=80, start=65, end=65)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        style_run(r, size=7, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    data = (
        ("H-07", "67.6%", "91.1%", "78.4%", "90.3%", "First-ranked holdout row under the toy rule"),
        ("H-18", "27.0%", "96.3%", "18.4%", "98.0%", "Peptide-free holdout comparator"),
    )
    for row_index, values in enumerate(data):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for col, value in enumerate(values):
            cells[col].width = Inches(widths[col])
            set_cell_margins(cells[col], top=75, bottom=75, start=60, end=60)
            set_cell_shading(cells[col], WHITE if row_index == 0 else GRAY_LIGHT)
            set_cell_border(cells[col], "CBD5E1", size="5")
            p = cells[col].paragraphs[0]
            if col < 5:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(value)
            style_run(r, size=7.3, bold=(col == 0), color=INK)


def add_figure(document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(paragraph, before=2, after=2)
    run = paragraph.add_run()
    inline = run.add_picture(str(FIGURE), width=Inches(6.55))
    doc_pr = inline._inline.docPr
    doc_pr.set("descr", "Predicted and frozen bacterial-signal reduction for six fully synthetic holdout formulations; circles are predictions and squares are frozen toy outcomes.")
    caption = document.add_paragraph(style="Caption")
    caption.paragraph_format.keep_with_next = True
    caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = caption.add_run(
        "Figure 1 | Predicted and frozen bacterial-signal reduction in six held-out synthetic formulations. "
        "Shape supplements color; no uncertainty is shown because independent replicate data do not exist."
    )
    style_run(r, size=8, color=GRAY)


def build() -> Path:
    if not FIGURE.exists():
        raise FileNotFoundError("Render figure-1.svg to figure-1.png before building the DOCX")

    document = Document()
    configure_document(document)
    configure_running_matter(document)

    add_banner(
        document,
        "Public before / after demonstration",
        "Same synthetic numbers.\nDefensible scientific claims.",
        "How CNS Skills turns fluent overclaiming into an evidence-bounded manuscript without inventing data.",
    )
    add_text(
        document,
        "FULLY SYNTHETIC · NO REAL RESEARCH · NO REAL CITATIONS · ALL 10.0000/cns.synthetic IDENTIFIERS ARE DELIBERATELY NON-RESOLVING",
        style="Demo Label",
        size=8,
        color=RED,
        bold=True,
        after=9,
    )
    add_metric_strip(document)
    add_heading(document, "What this artifact demonstrates", level=2)
    add_text(
        document,
        "The before manuscript uses polished language to outrun a fixed toy dataset. The after manuscript preserves every reported value while correcting autonomy, replication, safety, efficacy, mechanism, and translation claims. The objective is not cosmetic rewriting; it is visible control of the inference boundary.",
        size=10,
        after=8,
    )
    add_comparison_table(document)
    add_text(
        document,
        "Independent project. CNS means Cell · Nature · Science as an aspirational editorial benchmark; it is not publisher affiliation, endorsement, or an acceptance guarantee.",
        size=8.2,
        color=GRAY,
        italic=True,
        after=0,
    )

    add_page_break(document)
    add_kicker(document, "Before · intentionally flawed")
    add_heading(document, markdown_title(BEFORE), level=1, color=RED)
    add_text(document, "SYNTHETIC DEMONSTRATION — the language below is intentionally unsupported and must not be reused as scientific evidence.", style="Demo Label", size=8.3, color=RED, bold=True, after=7)
    for section_name in ("Abstract", "Methods", "Results", "Discussion"):
        add_heading(document, section_name, level=2, color=RED)
        section = markdown_section(BEFORE, section_name)
        if section_name == "Results":
            section = section.split("\n\nTable 1", 1)[0]
        add_manuscript_paragraphs(document, section)
    add_text(document, "Reader-visible leakage in the source also includes a TODO, an editorial assessment, an author query, and a ‘Revised version’ output label.", size=8.2, color=RED, bold=True, after=0)

    add_page_break(document)
    add_kicker(document, "Evidence boundary · decision record")
    add_heading(document, "The revision changes the inference—not the numbers", level=1)
    add_text(document, "Five gates convert the same fixture into a more defensible report. Automated checks expose structure; human scientific judgment decides whether a citation entails a claim and whether the stated validation layer is accurate.", size=9.6, after=8)
    add_risk_table(document)
    add_heading(document, "Source lock", level=2)
    add_text(document, "Frozen toy inputs: 18 fictional formulations; 12 training rows; six holdout rows; one fictional preparation batch; three emulated technical wells per endpoint; no random generator; no laboratory, animal, patient, or external publication data.", size=9.2)
    add_heading(document, "Permitted central claim", level=2)
    quote = document.add_table(rows=1, cols=1)
    quote.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = quote.cell(0, 0)
    set_cell_shading(cell, TEAL_LIGHT)
    set_cell_border(cell, TEAL, size="12", sides=("left",))
    set_cell_margins(cell, top=130, bottom=130, start=190, end=170)
    p = cell.paragraphs[0]
    r = p.add_run("A transparent ridge model retrospectively ranked six held-out synthetic rows under a toy viability constraint; this fixture demonstrates reporting discipline, not material discovery or biological performance.")
    style_run(r, size=10.2, bold=True, color=TEAL)
    add_heading(document, "What remains unresolved in a real study", level=2)
    add_text(document, "Independent synthesis batches, blinded outcome measurement, biological-replicate uncertainty, an appropriate antimicrobial comparator, prospective model-guided selection, mechanism testing, organism-level tolerability, and clinical relevance.", size=9.2, after=0)

    add_page_break(document)
    add_kicker(document, "After · evidence-bounded manuscript")
    add_heading(document, markdown_title(AFTER), level=1, color=TEAL)
    add_text(document, "FULLY SYNTHETIC — all values and placeholder references are invented for software demonstration.", style="Demo Label", size=8.3, color=RED, bold=True, after=7)
    add_heading(document, "Abstract", level=2, color=TEAL)
    add_manuscript_paragraphs(document, markdown_section(AFTER, "Abstract"))
    add_heading(document, "Methods", level=2, color=TEAL)
    add_manuscript_paragraphs(document, markdown_section(AFTER, "Methods"))

    add_page_break(document)
    add_kicker(document, "After · results and visual evidence")
    add_heading(document, "Results", level=1, color=TEAL)
    add_manuscript_paragraphs(document, markdown_section(AFTER, "Results"))
    add_figure(document)
    add_text(document, "Table 1 | Retrospective ranking outputs and frozen synthetic outcomes for H-07 and H-18.", style="Caption", size=8, color=GRAY, after=3)
    add_results_table(document)

    add_page_break(document)
    add_kicker(document, "After · discussion, citations, and audit state")
    add_heading(document, "Discussion", level=1, color=TEAL)
    add_manuscript_paragraphs(document, markdown_section(AFTER, "Discussion"))
    add_heading(document, "Synthetic placeholder references", level=2, color=TEAL)
    add_manuscript_paragraphs(document, markdown_section(AFTER, "References"))
    add_heading(document, "Validation snapshot", level=2, color=TEAL)
    add_text(document, "Deterministic fixture: PASS · strict clean-copy on revised source: PASS · display cross-references: clean · numeric citations: 3/3 placed · real or resolvable citations: 0", size=9, color=TEAL, bold=True, after=4)
    add_text(document, "These checks do not prove citation entailment, biological truth, or DOCX quality by themselves. This document was separately rendered page by page for visual QA before public use.", size=8.4, color=GRAY, italic=True, after=0)

    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
